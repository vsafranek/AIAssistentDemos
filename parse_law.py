import json
import re
from collections import deque
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os

# Get the current script's directory
SCRIPT_DIR = Path(__file__).resolve().parent

# Go two levels up to the project root
PROJECT_ROOT = SCRIPT_DIR.parents[0]


# ------------- File paths -------------
INPUT_DOCX = "./../Parser2/Sb_2000_219_2024-07-18_IZ.docx"
OUTPUT_JSON = "output_v5.json"
RULES_PATH = f"./rules.json"  # adjust if your configuration is elsewhere


def get_resource_path(relative_path):
    base_path = os.path.dirname(__file__)
    return os.path.join(base_path, relative_path)



# ------------- Load rules -------------

with open(get_resource_path('rules.json'), "r", encoding="utf-8") as f:
    RULES= json.load(f)


# ------------- Compile regexes from rules -------------
SECTION_RE = re.compile(RULES["section"]["keywords_regex"], re.IGNORECASE)
ARTICLE_RE = re.compile(RULES["article"]["detect_regex"], re.IGNORECASE)

NUM_RE = re.compile(RULES["numbering"]["regex"]["number"])
LET_RE = re.compile(RULES["numbering"]["regex"]["letter"])
ROM_RE = re.compile(RULES["numbering"]["regex"]["roman"], re.IGNORECASE)

# Article-internal paragraph number "(1) ..."
PARA_NUM_RE = re.compile(RULES["numbering"]["regex"].get("paragraph_number", r"^\((\d+)\)\s+"))
# Letter bullet "a) ..."
BULLET_LET_RE = re.compile(RULES["numbering"]["regex"].get("bullet_letter", r"^([A-Za-z])\)\s+"))
# Numeric subpoint "1." (allow both "1." and optionally "1)")
SUBPOINT_NUM_RE = re.compile(RULES["numbering"]["regex"].get("subpoint_number", r"^\s*(\d+)[\.)]\s+"))


# ------------- Formatting and property utilities -------------

def is_centered(para) -> bool:
    if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True
    try:
        st = para.style
        if st and st.paragraph_format and st.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
    except Exception:
        pass
    return False


def is_bold_like(para) -> bool:
    if any(bool(r.bold) for r in para.runs if r.text):
        return True
    try:
        st = para.style
        if st and getattr(st, "font", None) and getattr(st.font, "bold", None):
            return bool(st.font.bold)
    except Exception:
        pass
    return False


def get_rgb_hex(font):
    try:
        rgb = font.color.rgb if font and font.color else None
        return str(rgb).lower() if rgb else None
    except Exception:
        return None


def to_color_name(rgb_hex: str) -> str:
    if not rgb_hex:
        return RULES["color_map"].get("default", "black")
    rgb_hex = rgb_hex.lower()
    cmap = RULES["color_map"]
    if rgb_hex in cmap:
        return cmap[rgb_hex]

    tol = cmap.get("tolerance", {})
    if tol.get("use_hsv_buckets"):
        try:
            import colorsys
            if len(rgb_hex) == 6:
                r = int(rgb_hex[0:2], 16) / 255.0
                g = int(rgb_hex[2:4], 16) / 255.0
                b = int(rgb_hex[4:6], 16) / 255.0
                h, s, v = colorsys.rgb_to_hsv(r, g, b)
                h_deg = (h * 360.0) % 360.0

                def in_range(hd, lo, hi):
                    return lo <= hd <= hi

                def in_red(hd):
                    a, b = tol.get("red_hue_range", [345, 15])
                    return (hd >= a and hd <= 360) or (hd >= 0 and hd <= b)

                sat_min = tol.get("sat_min", 0.3)
                val_min = tol.get("val_min", 0.3)
                if s >= sat_min and v >= val_min:
                    if in_red(h_deg):
                        return "red"
                    lo_b, hi_b = tol.get("blue_hue_range", [200, 250])
                    if in_range(h_deg, lo_b, hi_b):
                        return "blue"
        except Exception:
            pass

    return cmap.get("default", "black")


def detect_double_strike(run) -> bool:
    try:
        el = run._element
        nsmap = dict(el.nsmap) if getattr(el, "nsmap", None) else {}
        nsmap.setdefault("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
        return bool(el.xpath(".//w:dstrike", namespaces=nsmap))
    except Exception:
        return False


def normalize_ws(s: str) -> str:
    if s is None:
        return ""
    txt = s
    norm = RULES.get("normalization", {})
    if norm.get("convert_nbsp_to_space", True):
        txt = txt.replace("\u00A0", " ").replace("\u202F", " ")
    if norm.get("replace_soft_break_with_space", True):
        txt = txt.replace("\v", " ").replace("\f", " ").replace("\u000B", " ")
    if norm.get("collapse_internal_spaces", False):
        txt = re.sub(r"[ \t]+", " ", txt)
    return txt


# ------------- Visual utilities -------------

def get_max_font_pt(para) -> float | None:
    mx = None
    for run in para.runs:
        try:
            sz = run.font.size
            if sz is not None:
                pt = sz.pt
                if mx is None or pt > mx:
                    mx = pt
        except Exception:
            continue
    if mx is None:
        try:
            st = para.style
            if st and getattr(st, "font", None) and getattr(st.font, "size", None):
                sz = st.font.size
                if sz is not None:
                    mx = sz.pt
        except Exception:
            pass
    return mx


def uppercase_ratio(s: str) -> float:
    letters = [ch for ch in s if ch.isalpha()]
    if not letters:
        return 0.0
    return sum(1 for ch in letters if ch.isupper()) / max(1, len(letters))


def looks_all_caps(s: str, threshold: float = 0.8) -> bool:
    return uppercase_ratio(s) >= threshold


def get_style_name(para) -> str | None:
    try:
        return para.style.name if para.style else None
    except Exception:
        return None


def get_para_visuals(para, text_norm: str) -> dict:
    return {
        "max_font_pt": get_max_font_pt(para),
        "is_bold": is_bold_like(para),
        "is_centered": is_centered(para),
        "upper_ratio": uppercase_ratio(text_norm),
        "is_all_caps_like": looks_all_caps(
            text_norm,
            threshold=RULES.get("visual", {}).get("all_caps_threshold", 0.8)
        ),
        "style_name": get_style_name(para),
    }


def estimate_body_font_pt(doc) -> float | None:
    sizes = []
    for p in doc.paragraphs:
        t = normalize_ws(p.text or "").strip()
        if not t:
            continue
        if not is_centered(p):
            pt = get_max_font_pt(p)
            if pt is not None:
                sizes.append(pt)
    if not sizes:
        return None
    sizes.sort()
    mid = len(sizes) // 2
    return sizes[mid] if len(sizes) % 2 == 1 else 0.5 * (sizes[mid - 1] + sizes[mid])


# ------------- Classification based on rule table -------------

def classify_by_rules(color: str, bold: bool, strike: bool, double_strike: bool) -> str:
    for rule in RULES["classification"]["rules"]:
        cond = rule["when"]
        if (cond.get("color") == color and
                cond.get("bold") == bool(bold) and
                cond.get("strike") == bool(strike) and
                cond.get("double_strike") == bool(double_strike)):
            return rule["label"]
    return RULES["classification"].get("fallback_label", "unknown")


def classify_visual_run(para, run, mode: str) -> tuple[str, dict]:
    rgb = get_rgb_hex(run.font)
    color = to_color_name(rgb)
    bold = bool(run.bold)
    strike = bool(getattr(run.font, "strike", False))
    double_strike = detect_double_strike(run)

    props = {
        "bold": bold,
        "strike": strike,
        "double_strike": double_strike,
        "color": color,
        "alignment": "center" if is_centered(para) else "other"
    }

    if mode == "heading_outside_article" and not RULES["heading_outside_article"].get("apply_rules", False):
        label = RULES["heading_outside_article"].get("label", "heading_text")
    else:
        label = classify_by_rules(color, bold, strike, double_strike)

    return label, props


def merge_key(label: str, props: dict) -> tuple:
    key_fields = RULES.get("merging", {}).get("merge_on_same", ["label", "bold", "strike", "double_strike", "color"])
    key = []
    for f in key_fields:
        if f == "label":
            key.append(label)
        else:
            key.append(props.get(f))
    return tuple(key)


# ------------- Paragraph segmentation -------------

def build_segments_from_paragraph(para, mode="text"):
    segments = []
    current = None

    for run in para.runs:
        raw = run.text or ""
        if raw == "":
            continue

        trimmed = normalize_ws(raw).strip()
        if not trimmed:
            if current is not None:
                current["text"] += raw
            continue

        label, props = classify_visual_run(para, run, mode=mode)
        key = merge_key(label, props)

        if current is None:
            current = {"label": label, "text": raw, "properties": props, "_k": key}
            continue

        if key == current["_k"]:
            current["text"] += raw
        else:
            segments.append({k: v for k, v in current.items() if k != "_k"})
            current = {"label": label, "text": raw, "properties": props, "_k": key}

    if current is not None:
        segments.append({k: v for k, v in current.items() if k != "_k"})

    cleaned = []
    for seg in segments:
        if normalize_ws(seg["text"]).strip():
            cleaned.append(seg)
    return cleaned


# ------------- Structure detection (part, article, point) -------------

def detect_part_heading(para_text: str, para) -> bool:
    txt = normalize_ws(para_text)
    m = SECTION_RE.search(txt)
    if not m:
        return False
    if RULES["section"].get("require_centered", True) and not is_centered(para):
        letters = [ch for ch in txt if ch.isalpha()]
        if letters:
            upper_ratio = sum(1 for ch in letters if ch.isupper()) / max(1, len(letters))
            return upper_ratio >= RULES["section"].get("uppercase_ratio_threshold", 0.7)
        return False
    return True


def detect_part_heading_visual(para_text: str, para, body_pt: float, part_font_mult: float) -> bool:
    txt = normalize_ws(para_text)
    vis = get_para_visuals(para, txt)
    if not vis["is_centered"]:
        return False
    mx = vis["max_font_pt"] or body_pt
    big_enough = mx >= body_pt * part_font_mult
    caps_ok = vis["is_all_caps_like"]
    by_regex = bool(SECTION_RE.search(txt))
    return by_regex or (big_enough and caps_ok and vis["is_bold"])


def detect_article_heading(para_text: str) -> bool:
    txt = normalize_ws(para_text)
    if RULES["article"].get("normalize_nbsp", True):
        txt = txt.replace("\u00A0", " ").replace("\u202F", " ")
    return bool(ARTICLE_RE.match(txt or ""))


def normalize_roman(s: str) -> str:
    return s.upper()


def classify_prefix(full_text: str):
    txt = normalize_ws(full_text or "")
    if not txt:
        return None

    m = NUM_RE.match(txt)
    if m:
        num_chain = m.group(1)
        rest = txt[m.end():]
        return {"type": "number", "key": num_chain, "rest_text": rest}

    m = LET_RE.match(txt)
    if m:
        letter = m.group(1)
        rest = txt[m.end():]
        return {"type": "letter", "key": letter.lower(), "rest_text": rest}

    m = ROM_RE.match(txt)
    if m:
        roman = normalize_roman(m.group(1))
        rest = txt[m.end():]
        return {"type": "roman", "key": roman, "rest_text": rest}

    return None


def compare_level_order(a_type: str, b_type: str) -> int:
    order_list = RULES["numbering"]["order"]
    idx_a = order_list.index(a_type)
    idx_b = order_list.index(b_type)
    return (idx_a > idx_b) - (idx_a < idx_b)


def count_number_depth(num_key: str) -> int:
    return 1 if not num_key else num_key.count(".") + 1


def is_same_level(prev_prefix, curr_prefix) -> bool:
    if not (prev_prefix and curr_prefix):
        return False
    if prev_prefix["type"] != curr_prefix["type"]:
        return False
    if prev_prefix["type"] == "number":
        return count_number_depth(prev_prefix["key"]) == count_number_depth(curr_prefix["key"])
    return True


def is_deeper(prev_prefix, curr_prefix) -> bool:
    if not (prev_prefix and curr_prefix):
        return False
    if prev_prefix["type"] == "number" and curr_prefix["type"] == "number":
        return count_number_depth(curr_prefix["key"]) > count_number_depth(prev_prefix["key"])
    return compare_level_order(prev_prefix["type"], curr_prefix["type"]) < 0


def is_shallower(prev_prefix, curr_prefix) -> bool:
    if not (prev_prefix and curr_prefix):
        return False
    if prev_prefix["type"] == "number" and curr_prefix["type"] == "number":
        return count_number_depth(curr_prefix["key"]) < count_number_depth(prev_prefix["key"])
    return compare_level_order(prev_prefix["type"], curr_prefix["type"]) > 0


# ------------- Node builders -------------

def make_node(node_type: str, title=None, meta=None):
    return {"type": node_type, "title": title, "meta": meta or {}, "children": []}


def make_paragraph_node(full_text: str, segments: list, prefix=None, para_heading=False):
    node_type = "heading" if para_heading else "paragraph"
    node = {
        "type": node_type,
        "title": normalize_ws(full_text).strip(),
        "meta": {
            "raw_text": full_text,
            "segments": segments
        },
        "children": []
    }
    if prefix:
        node["meta"]["prefix"] = {
            "type": prefix["type"],
            "key": prefix["key"]
        }
    return node


def aggregate_article_tags(heading_segments: list) -> dict:
    labels = sorted({seg["label"] for seg in heading_segments})
    colors = sorted({seg["properties"].get("color") for seg in heading_segments})
    has_bold = any(seg["properties"].get("bold") for seg in heading_segments)
    has_strike = any(seg["properties"].get("strike") for seg in heading_segments)
    return {
        "labels": labels,
        "colors": colors,
        "has_bold": has_bold,
        "has_strike": has_strike
    }


# ------------- Dedicated detectors for article hierarchy -------------

def classify_paragraph_prefix(full_text: str):
    txt = normalize_ws(full_text or "")
    m = PARA_NUM_RE.match(txt)
    if m:
        return {
            "type": "article_paragraph",
            "key": m.group(1),
            "rest_text": txt[m.end():]
        }
    return None


def classify_bullet_prefix(full_text: str):
    txt = normalize_ws(full_text or "")
    m = BULLET_LET_RE.match(txt)
    if m:
        return {
            "type": "point_letter",
            "key": m.group(1).lower(),
            "rest_text": txt[m.end():]
        }
    return None


def classify_subpoint_number(full_text: str):
    txt = normalize_ws(full_text or "")
    m = SUBPOINT_NUM_RE.match(txt)
    if m:
        return {
            "type": "subpoint_number",
            "key": m.group(1),
            "rest_text": txt[m.end():]
        }
    return None


# ------------- Visual-based heading helpers -------------

def is_visual_heading_candidate(para, txt_norm: str, body_pt: float, mult: float) -> bool:
    vis = get_para_visuals(para, txt_norm)
    if not vis["is_centered"] or not vis["is_bold"]:
        return False
    mx = vis["max_font_pt"] or body_pt
    return mx >= body_pt * mult


def try_compound_article_heading(paragraphs, i: int, body_pt: float, article_font_mult: float):
    if i + 1 >= len(paragraphs):
        return None
    A = paragraphs[i]
    B = paragraphs[i + 1]
    a_txt = normalize_ws(A.text or "").strip()
    b_txt = normalize_ws(B.text or "").strip()
    if not a_txt or not b_txt:
        return None

    if not is_visual_heading_candidate(A, a_txt, body_pt, article_font_mult):
        return None
    if ARTICLE_RE.match(a_txt):
        return None

    if not ARTICLE_RE.match(b_txt):
        return None
    if not is_centered(B):
        return None

    return {
        "A": A, "B": B,
        "A_text": a_txt, "B_text": b_txt
    }


def looks_like_inter_article_subtitle(para, txt_norm: str, body_pt: float, article_font_mult: float) -> bool:
    if not txt_norm:
        return False
    if ARTICLE_RE.match(txt_norm):
        return False
    if SECTION_RE.search(txt_norm):
        return False
    vis = get_para_visuals(para, txt_norm)
    if not vis["is_bold"]:
        return False
    mx = vis["max_font_pt"] or body_pt
    centered_or_big = vis["is_centered"] or (mx >= body_pt * (article_font_mult + 0.05))
    return centered_or_big and (mx >= body_pt * article_font_mult)


# ------------- Stack trimming helpers -------------

def trim_stack_to_level(stack: deque, level: str):
    """
    Normalize stack depth before inserting a new node of given level.
    - 'article_paragraph': keep only article_paragraph on top (pop deeper levels)
    - 'point_letter': pop subpoint levels; allow top to be article_paragraph or point_letter
    - 'subpoint_number': ensure top is point_letter if present, else article_paragraph (no pop if already point_letter)
    """
    def top_type():
        return stack[-1]["prefix"]["type"] if stack else None

    if level == "article_paragraph":
        while stack and top_type() != "article_paragraph":
            stack.pop()

    elif level == "point_letter":
        # pop all subpoints so that letters are direct children of article_paragraph
        while stack and top_type() == "subpoint_number":
            stack.pop()
        # if stack is deeper than point_letter (not typical), additional normalization not needed here

    elif level == "subpoint_number":
        # if we're at another subpoint, pop to its parent so next subpoint becomes sibling
        if stack and top_type() == "subpoint_number":
            stack.pop()
        # no further trimming; if top is article_paragraph (no letters used), we allow subpoint directly under it
        # if top is point_letter, it's ideal parent


# ------------- Main parsing -------------

def parse_doc_to_structure(doc_path: str) -> dict:
    doc = Document(doc_path)

    # Visual thresholds
    body_pt_est = estimate_body_font_pt(doc)
    default_body_pt = RULES.get("visual", {}).get("default_body_pt", 11.0)
    body_pt = body_pt_est or default_body_pt
    PART_FONT_MULT = RULES.get("visual", {}).get("part_font_multiplier", 1.5)
    ARTICLE_FONT_MULT = RULES.get("visual", {}).get("article_font_multiplier", 1.25)

    result = {
        "document": doc_path,
        "schema_version": RULES.get("schema_version", "1.0"),
        "parts": []
    }

    current_part = None
    current_article = None

    # stack entries: {"node": node, "prefix": {"type": "...", "key": "..."}}
    numbering_stack = deque()

    def reset_numbering():
        numbering_stack.clear()

    paras = list(doc.paragraphs)
    i = 0
    while i < len(paras):
        para = paras[i]
        full_text = para.text or ""
        stripped = normalize_ws(full_text).strip()
        if not stripped:
            i += 1
            continue

        # 0) Centered headings, parts and articles
        if is_centered(para) and RULES["section"].get("treat_centered_as_heading", True):
            # Part?
            if detect_part_heading(full_text, para) or detect_part_heading_visual(full_text, para, body_pt, PART_FONT_MULT):
                vis = get_para_visuals(para, stripped)
                current_part = make_node("part", title=stripped, meta={"raw_text": full_text, "visual": vis})
                result["parts"].append(current_part)
                current_article = None
                reset_numbering()
                segs = build_segments_from_paragraph(para, mode="heading_outside_article")
                current_part["children"].append(make_paragraph_node(full_text, segs, para_heading=True))
                i += 1
                continue

            # Two-line article heading?
            compound = try_compound_article_heading(paras, i, body_pt, ARTICLE_FONT_MULT)
            if compound:
                segA = build_segments_from_paragraph(compound["A"], mode="text")
                segB = build_segments_from_paragraph(compound["B"], mode="text")
                heading_segments = segA + segB
                article_tags = aggregate_article_tags(heading_segments)

                article_node = make_node(
                    "article",
                    title=compound["A_text"],
                    meta={
                        "raw_text": compound["A_text"] + "\n" + compound["B_text"],
                        "heading_segments": heading_segments,
                        "heading_summary": article_tags,
                        "article_number": compound["B_text"]
                    }
                )
                if current_part is None:
                    current_part = make_node("part", title=None)
                    result["parts"].append(current_part)
                current_part["children"].append(article_node)
                current_article = article_node
                reset_numbering()
                i += 2
                continue

            # Single-line article heading?
            if detect_article_heading(full_text):
                heading_segments = build_segments_from_paragraph(para, mode="text")
                article_tags = aggregate_article_tags(heading_segments)
                article_node = make_node(
                    "article",
                    title=None,
                    meta={
                        "raw_text": full_text,
                        "heading_segments": heading_segments,
                        "heading_summary": article_tags,
                        "article_number": stripped
                    }
                )
                if current_part is None:
                    current_part = make_node("part", title=None)
                    result["parts"].append(current_part)
                current_part["children"].append(article_node)
                current_article = article_node
                reset_numbering()
                i += 1
                continue

            # Other centered heading outside article
            if current_part is None:
                current_part = make_node("part", title=None)
                result["parts"].append(current_part)
            segs = build_segments_from_paragraph(para, mode="heading_outside_article")
            heading_node = make_paragraph_node(full_text, segs, para_heading=True)
            heading_node["meta"]["visual"] = get_para_visuals(para, stripped)
            current_part["children"].append(heading_node)
            i += 1
            continue

        # 1) Non-centered Part
        if detect_part_heading(full_text, para) or detect_part_heading_visual(full_text, para, body_pt, PART_FONT_MULT):
            vis = get_para_visuals(para, stripped)
            current_part = make_node("part", title=stripped, meta={"raw_text": full_text, "visual": vis})
            result["parts"].append(current_part)
            current_article = None
            reset_numbering()
            i += 1
            continue

        # 2) Non-centered Article
        if detect_article_heading(full_text):
            heading_segments = build_segments_from_paragraph(para, mode="text")
            article_tags = aggregate_article_tags(heading_segments)
            article_node = make_node(
                "article",
                title=None,
                meta={
                    "raw_text": full_text,
                    "heading_segments": heading_segments,
                    "heading_summary": article_tags,
                    "article_number": stripped
                }
            )
            if current_part is None:
                current_part = make_node("part", title=None)
                result["parts"].append(current_part)
            current_part["children"].append(article_node)
            current_article = article_node
            reset_numbering()
            i += 1
            continue

        # 2b) Inter-article subtitle before §
        if current_part is not None:
            if looks_like_inter_article_subtitle(para, stripped, body_pt, ARTICLE_FONT_MULT):
                j = i + 1
                next_nonempty = None
                while j < len(paras):
                    txtj = normalize_ws(paras[j].text or "").strip()
                    if txtj:
                        next_nonempty = txtj
                        break
                    j += 1
                if next_nonempty and detect_article_heading(next_nonempty):
                    segs = build_segments_from_paragraph(para, mode="heading_outside_article")
                    subtitle_node = make_paragraph_node(full_text, segs, para_heading=True)
                    subtitle_node["meta"]["visual"] = get_para_visuals(para, stripped)
                    subtitle_node["meta"]["binds_next_article"] = True
                    current_part["children"].append(subtitle_node)
                    numbering_stack.clear()
                    i += 1
                    continue

        # 3) Regular paragraph logic with article hierarchy
        segs = build_segments_from_paragraph(para, mode="text")

        if current_article is None:
            if current_part is None:
                current_part = make_node("part", title=None)
                result["parts"].append(current_part)
            current_article = make_node("article", title=None)
            current_part["children"].append(current_article)
            reset_numbering()

        # (1) paragraph level: resets deeper levels
        pfx_par = classify_paragraph_prefix(full_text)
        if pfx_par:
            numbering_stack.clear()
            paragraph_node = make_node("article_paragraph", title=pfx_par["key"], meta={"prefix_type": "number"})
            paragraph_node["children"].append(
                make_paragraph_node(
                    pfx_par["rest_text"],
                    segs,
                    prefix={"type": "number", "key": pfx_par["key"]}
                )
            )
            current_article["children"].append(paragraph_node)
            numbering_stack.append({"node": paragraph_node, "prefix": {"type": "article_paragraph", "key": pfx_par["key"]}})
            i += 1
            continue

        # a) letter level: must hang off article_paragraph, not subpoint
        pfx_bull = classify_bullet_prefix(full_text)
        if pfx_bull and numbering_stack:
            trim_stack_to_level(numbering_stack, "article_paragraph")
            if numbering_stack and numbering_stack[-1]["prefix"]["type"] == "article_paragraph":
                bullet_node = make_node("point", title=pfx_bull["key"], meta={"prefix_type": "letter"})
                bullet_node["children"].append(
                    make_paragraph_node(
                        pfx_bull["rest_text"],
                        segs,
                        prefix={"type": "letter", "key": pfx_bull["key"]}
                    )
                )
                numbering_stack[-1]["node"]["children"].append(bullet_node)
                # replace/append point_letter level; if previously a letter was open, we should replace it to keep same-level behavior
                # ensure top is point_letter
                if numbering_stack and numbering_stack[-1]["prefix"]["type"] == "point_letter":
                    numbering_stack.pop()
                numbering_stack.append({"node": bullet_node, "prefix": {"type": "point_letter", "key": pfx_bull["key"]}})
                i += 1
                continue

        # 1. subpoint: sibling under same letter (or directly under paragraph if letters absent)
        pfx_sub = classify_subpoint_number(full_text)
        if pfx_sub and numbering_stack:
            # ensure we don't stay stuck inside last subpoint
            trim_stack_to_level(numbering_stack, "subpoint_number")
            # now normalize so parent is either point_letter or article_paragraph
            if numbering_stack and numbering_stack[-1]["prefix"]["type"] == "subpoint_number":
                numbering_stack.pop()
            # if top is deeper than desired, pop handled above; if top is point_letter or article_paragraph, ok
            if numbering_stack:
                parent_type = numbering_stack[-1]["prefix"]["type"]
                if parent_type in ("point_letter", "article_paragraph"):
                    sub_node = make_node("subpoint", title=pfx_sub["key"], meta={"prefix_type": "sub_number"})
                    sub_node["children"].append(
                        make_paragraph_node(
                            pfx_sub["rest_text"],
                            segs,
                            prefix={"type": "sub_number", "key": pfx_sub["key"]}
                        )
                    )
                    numbering_stack[-1]["node"]["children"].append(sub_node)
                    # push subpoint level (replace previous subpoint if existed)
                    numbering_stack.append({"node": sub_node, "prefix": {"type": "subpoint_number", "key": pfx_sub["key"]}})
                    i += 1
                    continue

        # Plain paragraph: attach to deepest open node; do not change stack
        parent = current_article
        if numbering_stack:
            parent = numbering_stack[-1]["node"]
        parent["children"].append(make_paragraph_node(full_text, segs))
        i += 1

    return result


# ------------- Entrypoint -------------
if __name__ == "__main__":
    if not Path(INPUT_DOCX).exists():
        print(f"Warning: file does not exist: {INPUT_DOCX}")
    data = parse_doc_to_structure(INPUT_DOCX)
    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"Done. Written to {OUTPUT_JSON}")
